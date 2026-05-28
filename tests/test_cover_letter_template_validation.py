"""Tests for cover letter template validation."""

from __future__ import annotations

import tempfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from app.services.cover_letter_template_service import (
    REQUIRED_BODY_TAG,
    resolve_cover_letter_template_path,
    user_cover_letter_template_ready_for_build,
    validate_cover_letter_template_docx,
)
from app.services.resume_builder_service import fill_cover_letter_template


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def test_validate_requires_cover_letter_body_tag():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "cl.docx"
        _write_docx(path, ["Dear Hiring Manager,", "Zeyu Wang"])
        errors, _warnings, detected = validate_cover_letter_template_docx(path)
        assert any(REQUIRED_BODY_TAG in e for e in errors)
        assert "Zeyu Wang" not in detected


def test_validate_passes_with_body_tag_only():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "cl.docx"
        _write_docx(path, ["Zeyu Wang", "Dear Hiring Manager,", REQUIRED_BODY_TAG, "Sincerely,", "Zeyu Wang"])
        errors, warnings, detected = validate_cover_letter_template_docx(path)
        assert not errors
        assert REQUIRED_BODY_TAG in detected
        assert not warnings


def test_validate_warns_on_extra_placeholders():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "cl.docx"
        _write_docx(path, ["{{FULL_NAME}}", REQUIRED_BODY_TAG])
        errors, warnings, detected = validate_cover_letter_template_docx(path)
        assert not errors
        assert "{{FULL_NAME}}" in detected
        assert any("Only {{COVER_LETTER_BODY}} is filled automatically" in w for w in warnings)


def test_fill_cover_letter_template_only_replaces_body():
    with tempfile.TemporaryDirectory() as tmp:
        template_path = Path(tmp) / "template.docx"
        output_path = Path(tmp) / "filled.docx"
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "Zeyu Wang | Senior Software Engineer"
        doc.add_paragraph("Dear Hiring Manager,")
        doc.add_paragraph(REQUIRED_BODY_TAG)
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("Zeyu Wang")
        doc.save(str(template_path))

        fill_cover_letter_template(
            template_path,
            output_path,
            "First paragraph.\n\nSecond paragraph.",
        )
        assert output_path.exists()
        filled = Document(str(output_path))
        header_text = filled.sections[0].header.paragraphs[0].text
        body_text = "\n".join(p.text for p in filled.paragraphs)

        assert header_text == "Zeyu Wang | Senior Software Engineer"
        assert "First paragraph." in body_text
        assert "Second paragraph." in body_text
        assert REQUIRED_BODY_TAG not in body_text


def test_fill_cover_letter_template_handles_tag_split_across_runs():
    with tempfile.TemporaryDirectory() as tmp:
        template_path = Path(tmp) / "split.docx"
        output_path = Path(tmp) / "filled.docx"
        doc = Document()
        doc.add_paragraph("PROFILE")
        p = doc.add_paragraph()
        p.add_run("{{COVER_")
        p.add_run("LETTER_BODY}}")
        doc.save(str(template_path))

        fill_cover_letter_template(
            template_path,
            output_path,
            "Role-specific paragraph one.\n\nRole-specific paragraph two.",
        )
        filled = Document(str(output_path))
        body_text = "\n".join(p.text for p in filled.paragraphs)
        assert REQUIRED_BODY_TAG not in body_text
        assert "Role-specific paragraph one." in body_text
        assert "Role-specific paragraph two." in body_text


def test_fill_cover_letter_template_inherits_anchor_font():
    from docx.shared import Pt

    with tempfile.TemporaryDirectory() as tmp:
        template_path = Path(tmp) / "fonted.docx"
        output_path = Path(tmp) / "filled.docx"
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run(REQUIRED_BODY_TAG)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        doc.save(str(template_path))

        fill_cover_letter_template(
            template_path,
            output_path,
            "Hi Hiring Manager,\n\nMain body paragraph.\n\nBest regards,\nJane Doe",
        )
        filled = Document(str(output_path))
        body_paragraphs = [p for p in filled.paragraphs if p.text.strip()]
        assert any(p.text == "Hi Hiring Manager," for p in body_paragraphs)
        assert any("Main body paragraph." in p.text for p in body_paragraphs)
        signoff = next(p for p in body_paragraphs if "Best regards" in p.text)
        assert "Jane Doe" in signoff.text
        for p in body_paragraphs:
            for run in p.runs:
                if not run.text.strip():
                    continue
                assert run.font.name == "Times New Roman", (
                    f"Run {run.text!r} dropped Times New Roman font."
                )


def test_fill_cover_letter_template_signoff_uses_soft_break():
    """Single \\n inside a body part should render as a Word soft break (w:br)
    rather than splitting into two paragraphs — so 'Best regards,' and the name
    stay in the same paragraph but on separate lines.
    """
    import zipfile

    with tempfile.TemporaryDirectory() as tmp:
        template_path = Path(tmp) / "signoff.docx"
        output_path = Path(tmp) / "filled.docx"
        doc = Document()
        doc.add_paragraph(REQUIRED_BODY_TAG)
        doc.save(str(template_path))

        fill_cover_letter_template(
            template_path,
            output_path,
            "Hi Hiring Manager,\n\nMain body.\n\nBest regards,\nJane Doe",
        )
        with zipfile.ZipFile(output_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "<w:br/>" in xml or "<w:br />" in xml
        filled = Document(str(output_path))
        signoff = next(p for p in filled.paragraphs if "Best regards" in p.text)
        assert "Jane Doe" in signoff.text


def test_resolve_cover_letter_template_path_requires_user_upload():
    assert resolve_cover_letter_template_path(None) is None
    assert not user_cover_letter_template_ready_for_build(None)

    missing_user = SimpleNamespace(
        cover_letter_template_status="missing",
        cover_letter_template_working_path=None,
    )
    assert resolve_cover_letter_template_path(missing_user) is None

    with tempfile.TemporaryDirectory() as tmp:
        working = Path(tmp) / "working.docx"
        _write_docx(working, [REQUIRED_BODY_TAG])
        ready_user = SimpleNamespace(
            cover_letter_template_status="ready",
            cover_letter_template_working_path=str(working),
        )
        assert user_cover_letter_template_ready_for_build(ready_user)
        assert resolve_cover_letter_template_path(ready_user) == working
