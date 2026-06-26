"""Compile a ResumeDesign into a styled cover letter .docx template.

Reuses the same theme tokens (typography, colors, margins, header alignment) as the
resume so the cover letter matches visually. The generated template carries a single
``{{COVER_LETTER_BODY}}`` placeholder - letterhead, greeting, and signature are fixed
text - which the existing ``fill_cover_letter_template`` pipeline fills per job.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.database import User
from app.models.resume_design_schemas import ResumeDesign
from app.services.resume_design_compiler import _hex_to_rgb, _profile_dict, _set_run, _strip_headers

COVER_LETTER_BODY_TAG = "{{COVER_LETTER_BODY}}"


def compile_cover_letter_design(design: ResumeDesign, user: User, out_path: Path) -> list[str]:
    """Build a styled cover letter .docx for *design* at *out_path*; return detected tags."""
    profile = _profile_dict(user)
    typo = design.typography
    align = WD_ALIGN_PARAGRAPH.CENTER if design.layout.header_align == "center" else WD_ALIGN_PARAGRAPH.LEFT

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(design.layout.m_top)
    section.bottom_margin = Pt(design.layout.m_bottom)
    section.left_margin = Pt(design.layout.m_left)
    section.right_margin = Pt(design.layout.m_right)

    normal = doc.styles["Normal"]
    normal.font.name = typo.font_family
    normal.font.size = Pt(typo.base_font_pt)
    normal.font.color.rgb = _hex_to_rgb(design.colors.text)

    # Letterhead - name + contact line.
    name_para = doc.add_paragraph()
    name_para.alignment = align
    name_para.paragraph_format.space_after = Pt(2)
    _set_run(
        name_para.add_run(profile.get("full_name") or "Your Name"),
        font=typo.font_family,
        size_pt=typo.base_font_pt * max(typo.name_scale * 0.8, 1.4),
        color=_hex_to_rgb(design.colors.heading),
        bold=True,
    )

    contact_bits = [b for b in (profile.get("email"), profile.get("phone"), profile.get("linkedin"), profile.get("github")) if b]
    if contact_bits:
        cp = doc.add_paragraph()
        cp.alignment = align
        cp.paragraph_format.space_after = Pt(10)
        _set_run(
            cp.add_run("  |  ".join(contact_bits)),
            font=typo.font_family,
            size_pt=typo.base_font_pt * 0.95,
            color=_hex_to_rgb(design.colors.muted),
        )

    def _line(text: str, *, bold: bool = False, space_after: float = 10.0, color: str | None = None) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = typo.line_spacing
        _set_run(
            p.add_run(text),
            font=typo.font_family,
            size_pt=typo.base_font_pt,
            color=_hex_to_rgb(color or design.colors.text),
            bold=bold,
        )

    _line("Dear Hiring Manager,", space_after=10)
    _line(COVER_LETTER_BODY_TAG, space_after=12)
    _line("Sincerely,", space_after=2)
    _line(profile.get("full_name") or "Your Name", bold=True, space_after=0, color=design.colors.heading)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _strip_headers(out_path)
    return [COVER_LETTER_BODY_TAG]
