"""Low-level DOCX body structure helpers for template analysis and rendering."""

from __future__ import annotations

import re
from copy import deepcopy
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

TAG_PATTERN = re.compile(r"\{\{[^}]+\}\}")


def paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def table_text(table: Table) -> str:
    parts: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = " ".join(paragraph_text(p) for p in cell.paragraphs if paragraph_text(p).strip())
            cells.append(cell_text.strip())
        line = " | ".join(c for c in cells if c)
        if line:
            parts.append(line)
    return "\n".join(parts)


def iter_body_blocks(doc: Document) -> list[dict[str, Any]]:
    """Return ordered body blocks with stable indices for LLM outline + slicing."""
    blocks: list[dict[str, Any]] = []
    body = doc.element.body
    for idx, child in enumerate(body):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(child, doc)
            text = paragraph_text(p).strip()
            blocks.append({"index": idx, "kind": "paragraph", "text": text})
        elif tag == "tbl":
            t = Table(child, doc)
            text = table_text(t).strip()
            blocks.append({"index": idx, "kind": "table", "text": text})
    return blocks


def build_document_outline(doc: Document, *, max_blocks: int = 200) -> str:
    lines: list[str] = []
    for block in iter_body_blocks(doc)[:max_blocks]:
        text = (block.get("text") or "").replace("\n", " / ")
        if len(text) > 240:
            text = text[:240] + "…"
        lines.append(f"[{block['index']}] ({block['kind']}) {text}")
    return "\n".join(lines)


def _paragraphs_in_block(container: Any) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for paragraph in container.paragraphs:
        paragraphs.append(paragraph)
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def iter_document_paragraphs(doc: Document) -> list[Paragraph]:
    """All paragraphs in body, tables, and section headers/footers."""
    paragraphs = _paragraphs_in_block(doc)
    for section in doc.sections:
        for attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            try:
                part = getattr(section, attr, None)
            except ValueError:
                continue
            if part is None:
                continue
            paragraphs.extend(_paragraphs_in_block(part))
    return paragraphs


def find_tags_in_document(doc: Document) -> list[str]:
    found: list[str] = []
    seen: set[str] = set()
    for paragraph in iter_document_paragraphs(doc):
        text = paragraph_text(paragraph)
        for match in TAG_PATTERN.findall(text):
            if match not in seen:
                seen.add(match)
                found.append(match)
    return found


def body_child_elements(doc: Document) -> list:
    return list(doc.element.body)


def slice_body_elements(doc: Document, start_index: int, end_index: int) -> list:
    body = doc.element.body
    children = list(body)
    if start_index < 0 or end_index >= len(children) or start_index > end_index:
        return []
    return children[start_index : end_index + 1]


def remove_body_range(doc: Document, start_index: int, end_index: int) -> None:
    body = doc.element.body
    children = list(body)
    for idx in range(end_index, start_index - 1, -1):
        if 0 <= idx < len(children):
            body.remove(children[idx])


def insert_cloned_blocks_after(doc: Document, anchor_index: int, elements: list) -> None:
    body = doc.element.body
    children = list(body)
    if not children:
        return
    anchor = children[min(max(anchor_index, 0), len(children) - 1)]
    cursor = anchor
    for el in elements:
        clone = deepcopy(el)
        cursor.addnext(clone)
        cursor = clone


def replace_tag_in_document(doc: Document, tag: str, value: str) -> int:
    """Replace all occurrences of *tag* in body, tables, and section headers/footers."""
    replacements = 0
    for paragraph in iter_document_paragraphs(doc):
        if tag not in paragraph_text(paragraph):
            continue
        for run in paragraph.runs:
            if tag in run.text:
                run.text = run.text.replace(tag, value)
                replacements += 1
    return replacements


def remove_unresolved_tags(doc: Document) -> list[str]:
    leftover: list[str] = []
    seen: set[str] = set()
    for block in iter_body_blocks(doc):
        for match in TAG_PATTERN.findall(block.get("text") or ""):
            if match not in seen:
                seen.add(match)
                leftover.append(match)
            replace_tag_in_document(doc, match, "")
    return leftover
