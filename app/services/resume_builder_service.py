"""
Resume & cover letter document builder.

Opens user-designed DOCX templates, replaces placeholder tags with AI-tailored
content, and converts to PDF via LibreOffice.

Placeholders recognised in the resume template:
  {{PROFILE_SUMMARY}}  - single paragraph replacement
  {{SKILLS_CONTENT}}   - replaced by N skill-category rows
  {{EXP_1}} … {{EXP_N}} - replaced by per-company experience blocks

Placeholders recognised in the cover letter template:
  {{COVER_LETTER_BODY}} - AI-generated body only; all letterhead, greeting, and signature
  are fixed text in the user's uploaded template.
"""

from __future__ import annotations

import platform
import re
import shutil
import subprocess
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree

from app.core.config import get_settings
from app.core.logging import get_logger
from app.services.docx_structure import iter_document_paragraphs
from app.utils.resume_text_format import parse_bold_markers

logger = get_logger(__name__)

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COVER_LETTER_BODY_TAG = "{{COVER_LETTER_BODY}}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _w(tag: str) -> str:
    return f"{{{WNS}}}{tag}"


# ── Low-level XML helpers ──────────────────────────────────────────────────

def _clone_pPr(source_p, target_p) -> None:
    """Deep-copy *all* paragraph properties (style, numbering, spacing, indent, …)
    from *source_p* XML element to *target_p* XML element."""
    src_pPr = source_p.find(qn("w:pPr"))
    if src_pPr is None:
        return
    tgt_pPr = target_p.find(qn("w:pPr"))
    if tgt_pPr is not None:
        target_p.remove(tgt_pPr)
    target_p.insert(0, deepcopy(src_pPr))


def _clone_rPr(source_r) -> OxmlElement | None:
    """Deep-copy run properties from a source run element (or None)."""
    rPr = source_r.find(qn("w:rPr"))
    if rPr is None:
        return None
    return deepcopy(rPr)


def _get_template_rPr(anchor_p) -> OxmlElement | None:
    """Get a representative rPr from the anchor paragraph (first run or pPr/rPr)."""
    for r in anchor_p.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr)
    pPr = anchor_p.find(qn("w:pPr"))
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr)
    return None


def _runs_from_marked_text(
    text: str,
    rPr_template: OxmlElement | None,
    *,
    prefix: str = "",
) -> list[OxmlElement]:
    """Build Word runs from text that may contain ``**bold**`` markers."""
    segments = parse_bold_markers(text or "")
    runs: list[OxmlElement] = []
    if prefix:
        runs.append(_make_run(prefix, rPr_template, bold=False))
    for segment, is_bold in segments:
        if not segment:
            continue
        runs.append(_make_run(segment, rPr_template, bold=is_bold))
    if not runs and prefix:
        runs.append(_make_run(prefix, rPr_template, bold=False))
    return runs


def _set_paragraph_marked_text(paragraph: Paragraph, text: str, *, tag: str | None = None) -> None:
    """Replace paragraph content with runs that honor ``**bold**`` markers."""
    p_xml = paragraph._p
    rPr_tpl = _get_template_rPr(p_xml)
    full = "".join(run.text for run in paragraph.runs)
    if tag:
        full = full.replace(tag, text)
    else:
        full = text

    pPr = p_xml.find(qn("w:pPr"))
    for child in list(p_xml):
        if child is not pPr:
            p_xml.remove(child)

    for run_el in _runs_from_marked_text(full, rPr_tpl):
        p_xml.append(run_el)


def _make_run(text: str, rPr_template: OxmlElement | None = None, bold: bool = False) -> OxmlElement:
    """Create a <w:r> element with text and optional formatting."""
    r = OxmlElement("w:r")
    if rPr_template is not None:
        new_rPr = deepcopy(rPr_template)
        if bold:
            if new_rPr.find(qn("w:b")) is None:
                new_rPr.insert(0, OxmlElement("w:b"))
            if new_rPr.find(qn("w:bCs")) is None:
                new_rPr.append(OxmlElement("w:bCs"))
        else:
            for tag_name in ("w:b", "w:bCs"):
                el = new_rPr.find(qn(tag_name))
                if el is not None:
                    new_rPr.remove(el)
        r.append(new_rPr)
    elif bold:
        new_rPr = OxmlElement("w:rPr")
        new_rPr.append(OxmlElement("w:b"))
        new_rPr.append(OxmlElement("w:bCs"))
        r.append(new_rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _make_paragraph_from_anchor(anchor_p, runs: list[OxmlElement]) -> OxmlElement:
    """Create a new <w:p> that inherits all pPr from *anchor_p* and contains *runs*."""
    new_p = OxmlElement("w:p")
    _clone_pPr(anchor_p, new_p)
    for r in runs:
        new_p.append(r)
    return new_p


def _make_empty_spacer_paragraph(anchor_p) -> OxmlElement:
    """Build an empty paragraph inheriting *anchor_p*'s pPr so it renders as a
    single blank line at body line-height. Used to add visual breathing room
    between sections (e.g., description ↔ Key Contributions)."""
    return _make_paragraph_from_anchor(anchor_p, [])


# ── Placeholder replacement logic ──────────────────────────────────────────

def _find_paragraph_with_tag(doc: Document, tag: str) -> Paragraph | None:
    """Find the first paragraph whose combined run text contains *tag*."""
    for p in iter_document_paragraphs(doc):
        full = "".join(run.text for run in p.runs)
        if tag in full:
            return p
    return None


def _replace_inline_tag(paragraph: Paragraph, tag: str, replacement: str) -> bool:
    """Replace *tag* within a paragraph's runs, keeping all other runs/breaks intact.

    Works correctly even when the tag appears in a single run alongside
    line-breaks and other text (like the cover letter template).
    """
    for run in paragraph.runs:
        if tag in run.text:
            run.text = run.text.replace(tag, replacement)
            return True
    return False


def _replace_tag_with_paragraphs(
    doc: Document,
    tag: str,
    new_paragraphs: list[OxmlElement],
    *,
    cleanup_anchor: bool = True,
) -> bool:
    """Replace the paragraph containing *tag* with a list of new <w:p> elements.

    The new paragraphs are inserted *after* the anchor; then the anchor is removed
    (unless cleanup_anchor is False - used when the anchor has other content to keep).
    """
    anchor = _find_paragraph_with_tag(doc, tag)
    if anchor is None:
        return False

    anchor_xml = anchor._p
    parent = anchor_xml.getparent()

    cursor = anchor_xml
    for new_p in new_paragraphs:
        cursor.addnext(new_p)
        cursor = new_p

    if cleanup_anchor:
        parent.remove(anchor_xml)
    return True


def _split_anchor_around_tag(
    doc: Document,
    tag: str,
    body_paragraphs: list[OxmlElement],
) -> bool:
    """Handle the case where *tag* is embedded in a paragraph alongside other content
    (e.g., the cover letter template: ``{{DATE}} <br> Dear Hiring Manager, <br> {{COVER_LETTER_BODY}}``).

    Strategy: find the run containing *tag*, remove the tag text from the run,
    and insert body paragraphs *after* the anchor paragraph. The anchor paragraph
    is preserved (it contains the date and greeting).
    """
    anchor = _find_paragraph_with_tag(doc, tag)
    if anchor is None:
        return False

    anchor_xml = anchor._p

    # Find the specific run that contains the tag
    tag_run = None
    tag_run_idx = None
    for idx, r in enumerate(anchor_xml.findall(qn("w:r"))):
        t_el = r.find(qn("w:t"))
        if t_el is not None and t_el.text and tag in t_el.text:
            tag_run = r
            tag_run_idx = idx
            break

    if tag_run is None:
        return False

    # Remove the tag text from the run
    t_el = tag_run.find(qn("w:t"))
    t_el.text = t_el.text.replace(tag, "")

    # Remove any trailing <w:br> elements after the tag run (they were separating
    # the tag from the next content, which is now gone)
    all_children = list(anchor_xml)
    pPr = anchor_xml.find(qn("w:pPr"))
    runs_and_brs = [c for c in all_children if c is not pPr] if pPr is not None else list(all_children)

    # Find elements after the tag run and remove trailing breaks/empty runs
    found_tag = False
    to_remove = []
    for child in runs_and_brs:
        if child is tag_run:
            found_tag = True
            # If the tag run is now empty, remove it too
            if not (t_el.text and t_el.text.strip()):
                to_remove.append(child)
            continue
        if found_tag:
            # Remove trailing breaks and empty runs after tag
            if child.tag == qn("w:r"):
                has_br = child.find(qn("w:br")) is not None
                has_text = False
                for sub_t in child.findall(qn("w:t")):
                    if sub_t.text and sub_t.text.strip():
                        has_text = True
                if has_br and not has_text:
                    to_remove.append(child)
                else:
                    break
            else:
                break

    for el in to_remove:
        anchor_xml.remove(el)

    # Insert body paragraphs after the anchor
    cursor = anchor_xml
    for new_p in body_paragraphs:
        cursor.addnext(new_p)
        cursor = new_p

    return True


# ── Resume content builders ───────────────────────────────────────────────

_SKILLS_SPLIT_RE = re.compile(r"[,;|\n]+")


def _split_skill_values(s: str) -> list[str]:
    return [x.strip() for x in _SKILLS_SPLIT_RE.split(s or "") if x.strip()]


def _tint_hex(value: str | None, keep: float) -> str:
    """Blend *value* toward white (keep=fraction of original retained)."""
    v = (value or "#000000").lstrip("#")
    if len(v) == 3:
        v = "".join(ch * 2 for ch in v)
    try:
        r, g, b = int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16)
    except Exception:
        return "f1f5f9"
    r = round(255 * (1 - keep) + r * keep)
    g = round(255 * (1 - keep) + g * keep)
    b = round(255 * (1 - keep) + b * keep)
    return f"{r:02x}{g:02x}{b:02x}"


def _styled_run(
    text: str,
    rPr_template: OxmlElement | None,
    *,
    bold: bool = False,
    caps: bool = False,
    color_hex: str | None = None,
    fill_hex: str | None = None,
) -> OxmlElement:
    """A <w:r> with optional bold, caps, font color and background shading."""
    r = OxmlElement("w:r")
    rPr = deepcopy(rPr_template) if rPr_template is not None else OxmlElement("w:rPr")
    for tag_name in ("w:b", "w:bCs", "w:caps", "w:color", "w:shd"):
        for el in rPr.findall(qn(tag_name)):
            rPr.remove(el)
    if bold:
        rPr.insert(0, OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if caps:
        rPr.append(OxmlElement("w:caps"))
    if color_hex:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), color_hex.lstrip("#"))
        rPr.append(col)
    if fill_hex:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#"))
        rPr.append(shd)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _category_runs(
    cat: str,
    rPr_tpl: OxmlElement | None,
    cat_style: str,
    accent: str | None,
    heading: str | None,
) -> list[OxmlElement]:
    if cat_style == "caps":
        return [_styled_run(cat.upper(), rPr_tpl, bold=True, color_hex=heading)]
    if cat_style == "accent":
        return [_styled_run(cat, rPr_tpl, bold=True, color_hex=accent or heading)]
    if cat_style == "badge":
        return [_styled_run(f"  {cat.upper()}  ", rPr_tpl, bold=True, color_hex="#ffffff", fill_hex=accent or "334155")]
    if cat_style == "bar":
        runs: list[OxmlElement] = []
        runs.append(_styled_run("\u258f ", rPr_tpl, bold=True, color_hex=accent or heading))
        runs.append(_styled_run(cat, rPr_tpl, bold=True, color_hex=heading))
        return runs
    return [_styled_run(cat, rPr_tpl, bold=True, color_hex=heading)]


def _build_skills_elements(
    skills: list[dict],
    anchor_p,
    style: dict | None = None,
    colors: dict | None = None,
) -> list[OxmlElement]:
    """Build <w:p> elements for the skills section, cloning formatting from anchor.

    Each skill in a category is split into its own structured unit; the SkillsStyle
    *style* controls the layout (inline / stacked / bullets / chips / pipe / grid),
    category-label treatment and chip shading so the .docx matches the live preview.
    """
    rPr_tpl = _get_template_rPr(anchor_p)
    style = style or {}
    colors = colors or {}
    layout = style.get("layout", "inline")
    if layout == "bullets":  # one-term-per-line lists were dropped; render as chips
        layout = "chips"
    cat_style = style.get("category", "bold")
    accent_chips = bool(style.get("accent_chips"))
    accent = (colors.get("accent") or "").strip() or None
    heading = (colors.get("heading") or "").strip() or None
    text_col = (colors.get("text") or "").strip() or None

    result: list[OxmlElement] = []
    for item in skills:
        cat = (item.get("category") or "").strip()
        vals = _split_skill_values(item.get("skills") or "")
        if not vals and not cat:
            continue
        cat_runs = _category_runs(cat, rPr_tpl, cat_style, accent, heading) if cat else []

        if layout in ("inline", "pipe", "grid"):
            sep = "  |  " if layout == "pipe" else ", "
            runs = list(cat_runs)
            if cat:
                runs.append(_make_run(" " if cat_style == "badge" else ": ", rPr_tpl))
            runs.extend(_runs_from_marked_text(sep.join(vals), rPr_tpl))
            result.append(_make_paragraph_from_anchor(anchor_p, runs))
        elif layout == "stacked":
            if cat:
                result.append(_make_paragraph_from_anchor(anchor_p, cat_runs))
            result.append(_make_paragraph_from_anchor(anchor_p, _runs_from_marked_text(", ".join(vals), rPr_tpl)))
        elif layout == "chips":
            if cat:
                result.append(_make_paragraph_from_anchor(anchor_p, cat_runs))
            fill = _tint_hex(accent, 0.14) if (accent_chips and accent) else "eef2f7"
            chip_color = accent if (accent_chips and accent) else text_col
            chip_runs: list[OxmlElement] = []
            for i, sk in enumerate(vals):
                if i:
                    chip_runs.append(_make_run("  ", rPr_tpl))
                chip_runs.append(_styled_run(f"  {sk}  ", rPr_tpl, color_hex=chip_color, fill_hex=fill))
            result.append(_make_paragraph_from_anchor(anchor_p, chip_runs))
        else:
            runs = list(cat_runs)
            if cat:
                runs.append(_make_run(": ", rPr_tpl))
            runs.extend(_runs_from_marked_text(", ".join(vals), rPr_tpl))
            result.append(_make_paragraph_from_anchor(anchor_p, runs))
    return result


_INLINE_BULLET_RE = re.compile(r"[•▪‣◦∙·●]")
_LINE_BULLET_RE = re.compile(r"^\s*(?:[-*▪‣◦∙·●]|\d+[.)])\s+(.*)$")
_PROJECT_LINE_RE = re.compile(r"^\s*project\s*[:\-\u2013\u2014]\s*", re.IGNORECASE)


def _build_lead_paragraphs(
    lead: str,
    anchor_p,
    rPr_tpl,
    *,
    project_already_rendered: bool,
) -> list[OxmlElement]:
    """Render the experience lead text, putting a ``Project: <title>`` prefix on its
    own line (bold label) and keeping the remaining description as a flowing paragraph.

    The project title and description frequently arrive separated by a newline that
    would otherwise be collapsed into one run-on line; this restores the line break.
    """
    segments = [seg.strip() for seg in (lead or "").split("\n") if seg.strip()]
    if not segments:
        return []

    paragraphs: list[OxmlElement] = []
    start = 0
    first = segments[0]
    looks_like_title = bool(_PROJECT_LINE_RE.match(first)) and (len(segments) > 1 or len(first) <= 80)

    if looks_like_title:
        if project_already_rendered:
            # A PROJECT: line was already emitted from project_name - drop the duplicate.
            start = 1
        else:
            match = _PROJECT_LINE_RE.match(first)
            title = first[match.end():].strip()
            runs = [_make_run("Project: ", rPr_tpl, bold=True)]
            if title:
                runs.extend(_runs_from_marked_text(title, rPr_tpl))
            paragraphs.append(_make_paragraph_from_anchor(anchor_p, runs))
            start = 1

    description = " ".join(segments[start:]).strip()
    if description:
        paragraphs.append(_make_paragraph_from_anchor(anchor_p, _runs_from_marked_text(description, rPr_tpl)))
    return paragraphs


def split_description_and_bullets(text: str) -> tuple[str, list[str]]:
    """Separate a free-form experience description into a lead paragraph and bullets.

    Tailored/profile text frequently arrives as a single blob where individual
    achievements are joined inline with a bullet glyph (``… workloads. • Architected …
    • Built …``) or as newline-prefixed list items. Rendering that verbatim produces
    one giant run-on paragraph. This splits it so each achievement becomes its own
    bullet line while keeping any introductory sentence as the lead.
    """
    s = (text or "").strip()
    if not s:
        return "", []

    if _INLINE_BULLET_RE.search(s):
        segments = [seg.strip(" \t\r\n-\u2013\u2014").strip() for seg in _INLINE_BULLET_RE.split(s)]
        segments = [seg for seg in segments if seg]
        if len(segments) >= 2:
            return segments[0], segments[1:]
        return s, []

    lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
    matches = [_LINE_BULLET_RE.match(ln) for ln in lines]
    if sum(1 for m in matches if m) >= 2:
        lead_lines: list[str] = []
        bullets: list[str] = []
        for ln, m in zip(lines, matches):
            if m:
                bullets.append(m.group(1).strip())
            elif not bullets:
                lead_lines.append(ln)
        return " ".join(lead_lines).strip(), bullets

    return s, []


_EXP_MARKER_GLYPH: dict[str, str] = {
    "dot": "\u2022",
    "dash": "\u2013",
    "arrow": "\u2192",
    "chevron": "\u203A",
    "square": "\u25AA",
    "diamond": "\u25C6",
    "none": "",
}


def _exp_label_runs(label_style: str, rPr_tpl, accent: str | None, heading: str | None) -> list[OxmlElement]:
    text = "Key Contributions:"
    if label_style == "bold":
        return [_styled_run(text, rPr_tpl, bold=True, color_hex=heading)]
    if label_style == "accent":
        return [_styled_run(text, rPr_tpl, bold=True, color_hex=accent or heading)]
    if label_style == "caps":
        return [_styled_run("KEY CONTRIBUTIONS:", rPr_tpl, bold=True, color_hex=heading)]
    return [_make_run(text, rPr_tpl, bold=False)]


def _exp_used_skills_paragraphs(
    used_skills: str,
    used_style: str,
    anchor_p,
    rPr_tpl,
    accent: str | None,
    heading: str | None,
    text_col: str | None,
) -> list[OxmlElement]:
    used_skills = (used_skills or "").strip()
    if not used_skills:
        return []
    if used_style == "inline":
        runs = [_styled_run("Technologies: ", rPr_tpl, bold=True, color_hex=heading)]
        runs.extend(_runs_from_marked_text(used_skills, rPr_tpl))
        return [_make_paragraph_from_anchor(anchor_p, runs)]
    if used_style == "label":
        runs = [_styled_run("Tech \u00b7 ", rPr_tpl, bold=True, color_hex=accent or heading)]
        runs.extend(_runs_from_marked_text(used_skills, rPr_tpl))
        return [_make_paragraph_from_anchor(anchor_p, runs)]
    # chips / pill - each skill is a shaded segment
    accent_pill = used_style == "pill"
    fill = _tint_hex(accent, 0.16) if (accent_pill and accent) else "eef2f7"
    chip_color = accent if (accent_pill and accent) else text_col
    runs: list[OxmlElement] = []
    for sk in _split_skill_values(used_skills):
        runs.append(_styled_run(f"  {sk}  ", rPr_tpl, color_hex=chip_color, fill_hex=fill))
        runs.append(_make_run("  ", rPr_tpl))
    if not runs:
        return []
    return [_make_paragraph_from_anchor(anchor_p, runs)]


def _build_experience_body(
    exp: dict,
    anchor_p,
    project_header_p,
    style: dict | None,
    colors: dict | None,
) -> list[OxmlElement]:
    """Build the body <w:p> elements for one company's experience block: project
    title, intro, 'Key Contributions:' label, contribution bullets and used skills.

    The company / role / date header is rendered separately (at template-compile time
    for builder themes, or carried by the user's own template), so this never emits a
    header. *style* is the ExperienceStyle dict (control board + item styling)."""
    rPr_tpl = _get_template_rPr(anchor_p)
    header_p_ref = project_header_p if project_header_p is not None else anchor_p
    style = style or {}
    colors = colors or {}
    accent = (colors.get("accent") or "").strip() or None
    heading = (colors.get("heading") or "").strip() or None
    text_col = (colors.get("text") or "").strip() or None

    project_style = style.get("project_style", "label")
    show_project = style.get("show_project_title", True)
    intro_style = style.get("intro_style", "plain")
    show_intro = style.get("show_intro", True)
    marker = style.get("marker", "dot")
    label_style = style.get("label_style", "plain")
    show_label = style.get("show_contributions_label", True)
    used_style = style.get("used_skills_style", "inline")
    show_used = style.get("show_used_skills", True)

    paragraphs: list[OxmlElement] = []

    project_name = exp.get("project_name")
    project_rendered = (
        bool(project_name) and project_name not in ("None", "null") and show_project and project_style != "hidden"
    )
    if project_rendered:
        header_rPr = _get_template_rPr(header_p_ref)
        p = OxmlElement("w:p")
        _clone_pPr(header_p_ref, p)
        if project_style == "label":
            p.append(_make_run("PROJECT:", header_rPr, bold=True))
            p.append(_make_run(f" {project_name}", header_rPr, bold=False))
        elif project_style == "accent":
            p.append(_styled_run(project_name, header_rPr, bold=True, color_hex=accent or heading))
        elif project_style == "bold":
            p.append(_styled_run(project_name, header_rPr, bold=True, color_hex=heading))
        else:  # italic degrades to plain in DOCX
            p.append(_make_run(project_name, header_rPr))
        paragraphs.append(p)

    lead, inline_bullets = split_description_and_bullets(exp.get("project_description", ""))
    if show_intro and intro_style != "hidden":
        lead_paragraphs = _build_lead_paragraphs(
            lead, anchor_p, rPr_tpl, project_already_rendered=project_rendered
        )
    else:
        lead_paragraphs = []
    paragraphs.extend(lead_paragraphs)

    explicit_bullets = [str(b) for b in (exp.get("bullets") or []) if str(b).strip()]
    bullets = [*inline_bullets, *explicit_bullets]
    if bullets:
        if lead_paragraphs:
            paragraphs.append(_make_empty_spacer_paragraph(anchor_p))

        if show_label and label_style != "hidden":
            paragraphs.append(
                _make_paragraph_from_anchor(anchor_p, _exp_label_runs(label_style, rPr_tpl, accent, heading))
            )

        for idx, bullet_text in enumerate(bullets, start=1):
            runs: list[OxmlElement] = []
            if marker == "numbered":
                runs.append(_styled_run(f"{idx}. ", rPr_tpl, bold=True, color_hex=accent))
            else:
                glyph = _EXP_MARKER_GLYPH.get(marker, "\u2022")
                if glyph:
                    runs.append(_styled_run(f"{glyph} ", rPr_tpl, bold=True, color_hex=accent))
            runs.extend(_runs_from_marked_text(str(bullet_text), rPr_tpl))
            paragraphs.append(_make_paragraph_from_anchor(anchor_p, runs))

    if show_used and used_style != "hidden":
        paragraphs.extend(
            _exp_used_skills_paragraphs(
                exp.get("used_skills", ""), used_style, anchor_p, rPr_tpl, accent, heading, text_col
            )
        )

    return paragraphs


def _build_experience_elements(
    exp: dict,
    anchor_p,
    project_header_p=None,
    style: dict | None = None,
    colors: dict | None = None,
) -> list[OxmlElement]:
    """Style-aware body builder for one company's ``{{EXP_N}}`` block.

    Kept as a thin wrapper around :func:`_build_experience_body` for backward
    compatibility with callers that do not pass a style.
    """
    return _build_experience_body(exp, anchor_p, project_header_p, style, colors)


def _clean_leftover_exp_placeholders(doc: Document) -> None:
    """Clear any unreplaced {{EXP_N}} placeholder text (replace with empty).

    We keep the paragraph element so the surrounding company header tables and
    spacing are not disrupted - only the tag text itself is removed.
    """
    pattern = re.compile(r"\{\{EXP_\d+\}\}")
    for p in doc.paragraphs:
        for run in p.runs:
            if pattern.search(run.text):
                run.text = pattern.sub("", run.text)


# ── Layout-preserving DOCX serialization ───────────────────────────────────

RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
HEADER_REL_TYPE = f"{OFFICE_REL_NS}/header"
FOOTER_REL_TYPE = f"{OFFICE_REL_NS}/footer"


def _read_template_header_footer_rids(template_path: Path) -> set[str]:
    """Return the set of rIds in the template that legitimately point to a
    real ``word/header*.xml`` or ``word/footer*.xml`` part. Any other
    header/footer reference in the modified output is a phantom one injected
    by python-docx that must be stripped to avoid layout shifts.
    """
    try:
        with zipfile.ZipFile(template_path, "r") as zf:
            try:
                rels_xml = zf.read("word/_rels/document.xml.rels")
            except KeyError:
                return set()
    except zipfile.BadZipFile:
        return set()

    root = etree.fromstring(rels_xml)
    rids: set[str] = set()
    for rel in root.findall(f"{{{RELS_NS}}}Relationship"):
        rel_type = rel.get("Type")
        if rel_type in (HEADER_REL_TYPE, FOOTER_REL_TYPE):
            rid = rel.get("Id")
            if rid:
                rids.add(rid)
    return rids


def _strip_phantom_header_footer_refs(root: etree._Element, allowed_rids: set[str]) -> None:
    """Remove ``<w:headerReference>`` / ``<w:footerReference>`` entries whose
    ``r:id`` is not in *allowed_rids*. These are the phantom references
    python-docx injects on save - keeping only ones the user truly authored.
    """
    rid_attr = f"{{{OFFICE_REL_NS}}}id"
    for ref_tag in (_w("headerReference"), _w("footerReference")):
        for ref in root.iter(ref_tag):
            rid = ref.get(rid_attr)
            if rid not in allowed_rids:
                parent = ref.getparent()
                if parent is not None:
                    parent.remove(ref)


def _save_docx_preserving_template_layout(
    doc: Document,
    template_path: Path,
    output_path: Path,
) -> None:
    """Write *doc* to *output_path* without inheriting python-docx's package mutations.

    ``python-docx`` rewrites the full DOCX package on ``doc.save()`` and
    auto-injects empty ``<w:headerReference>``/``<w:footerReference>`` entries
    into every ``<w:sectPr>`` plus phantom ``word/header*.xml`` /
    ``word/footer*.xml`` files into the archive. Word/LibreOffice then
    reserves header/footer space on every page even though those files are
    empty - pushing the body content down and visibly shifting the user's
    designed layout.

    Strategy:
      1. Serialize the in-memory document (with all placeholder replacements)
         using lxml so we keep the modified body content but skip
         ``doc.save()`` entirely (avoiding package-level mutations).
      2. Strip any ``<w:headerReference>`` / ``<w:footerReference>`` whose
         ``r:id`` does NOT correspond to a real header/footer relationship in
         the original template's ``word/_rels/document.xml.rels``. Genuine
         user-authored references are preserved.
      3. Copy the original template file byte-for-byte to *output_path*
         (preserves ``[Content_Types].xml``, all ``word/_rels``, fonts,
         styles, images, theme, custom XML, and any real header/footer files).
      4. Replace ONLY ``word/document.xml`` inside that zip with the cleaned
         XML so no package-level metadata or phantom parts leak through.
    """
    new_root_xml = etree.tostring(doc.element)
    new_root = etree.fromstring(new_root_xml)
    if new_root.find(_w("body")) is None:
        raise RuntimeError("Modified resume document is missing <w:body>.")

    allowed_rids = _read_template_header_footer_rids(template_path)
    _strip_phantom_header_footer_refs(new_root, allowed_rids)

    final_xml = etree.tostring(
        new_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    shutil.copy2(template_path, output_path)
    _replace_docx_internal(output_path, "word/document.xml", final_xml)


# ── Public API ─────────────────────────────────────────────────────────────

def fill_resume_template(
    template_path: Path,
    output_path: Path,
    tailored: dict[str, Any],
) -> Path:
    """Fill resume template with tailored content and save to *output_path*."""
    doc = Document(str(template_path))

    # Profile summary - simple inline replacement
    summary = tailored.get("profile_summary", "")
    anchor = _find_paragraph_with_tag(doc, "{{PROFILE_SUMMARY}}")
    if anchor:
        _set_paragraph_marked_text(anchor, summary, tag="{{PROFILE_SUMMARY}}")

    # Technical skills - replace with multiple skill-category paragraphs
    skills = tailored.get("technical_skills", [])
    if skills:
        skills_anchor = _find_paragraph_with_tag(doc, "{{SKILLS_CONTENT}}")
        if skills_anchor:
            elements = _build_skills_elements(
                skills, skills_anchor._p, tailored.get("skills_style"), tailored.get("colors")
            )
            _replace_tag_with_paragraphs(doc, "{{SKILLS_CONTENT}}", elements)

    # Capture the PROJECT: header paragraph XML as a formatting template.
    # The template has a fixed "PROJECT: ..." line before {{EXP_1}} - we clone
    # its formatting for all other companies' PROJECT headers.
    project_header_ref = None
    for p in doc.paragraphs:
        full = "".join(run.text for run in p.runs)
        if full.strip().startswith("PROJECT:"):
            project_header_ref = p._p
            break

    # Work experience - replace each {{EXP_N}} placeholder
    experience = tailored.get("work_experience", [])
    exp_style = tailored.get("experience_style")
    exp_colors = tailored.get("colors")
    for idx, exp in enumerate(experience, start=1):
        tag = "{{" + f"EXP_{idx}" + "}}"
        exp_anchor = _find_paragraph_with_tag(doc, tag)
        if exp_anchor:
            if idx == 1 and project_header_ref is not None:
                # EXP_1: the template already has a PROJECT: line - update it with the
                # AI project name, then fill the body without re-rendering the project.
                project_name = exp.get("project_name")
                if project_name and project_name not in ("None", "null"):
                    for run in Paragraph(project_header_ref, None).runs:
                        full = run.text
                        if full.strip().startswith("PROJECT:"):
                            continue
                        if full.strip() and not full.strip().startswith("PROJECT"):
                            run.text = project_name
                            break

                # Project header is owned by the template here, so suppress it in the body.
                body_style = {**(exp_style or {}), "show_project_title": False}
                body_elements = _build_experience_body(
                    exp, exp_anchor._p, None, body_style, exp_colors
                )
                if body_elements:
                    _replace_tag_with_paragraphs(doc, tag, body_elements)
                else:
                    _replace_inline_tag(exp_anchor, tag, "")
            else:
                elements = _build_experience_elements(
                    exp, exp_anchor._p, project_header_ref, exp_style, exp_colors
                )
                if elements:
                    _replace_tag_with_paragraphs(doc, tag, elements)
                else:
                    _replace_inline_tag(exp_anchor, tag, "")

    _clean_leftover_exp_placeholders(doc)

    _save_docx_preserving_template_layout(doc, template_path, output_path)
    logger.info("resume_docx_created", path=str(output_path))
    return output_path


def _paragraph_text_from_xml(p_el: etree._Element) -> str:
    return "".join(t.text or "" for t in p_el.iter(_w("t")))


def _iter_document_body_paragraphs(body_el: etree._Element):
    for child in body_el:
        local = child.tag.split("}")[-1]
        if local == "p":
            yield child
        elif local == "tbl":
            for tc in child.iter(_w("tc")):
                for p_el in tc.findall(_w("p")):
                    yield p_el


def _extract_anchor_rpr(p_el: etree._Element) -> etree._Element | None:
    """Return a deep copy of the most representative ``<w:rPr>`` for new runs.

    Preference order:
      1. The ``<w:rPr>`` of the first run that contains visible text in *p_el*.
      2. The ``<w:rPr>`` of any run.
      3. The paragraph-level default run formatting at ``<w:pPr>/<w:rPr>``.
    """
    runs = p_el.findall(_w("r"))
    for r in runs:
        if r.find(_w("t")) is not None:
            rpr = r.find(_w("rPr"))
            if rpr is not None:
                return deepcopy(rpr)
    for r in runs:
        rpr = r.find(_w("rPr"))
        if rpr is not None:
            return deepcopy(rpr)
    p_pr = p_el.find(_w("pPr"))
    if p_pr is not None:
        rpr = p_pr.find(_w("rPr"))
        if rpr is not None:
            return deepcopy(rpr)
    return None


def _clear_paragraph_content(p_el: etree._Element) -> etree._Element | None:
    p_pr = p_el.find(_w("pPr"))
    for child in list(p_el):
        if child is not p_pr:
            p_el.remove(child)
    return p_pr


def _append_text_run(
    p_el: etree._Element,
    text: str,
    rpr_template: etree._Element | None,
) -> None:
    """Append a single ``<w:r><w:t>`` to *p_el* preserving font formatting."""
    if not text:
        return
    run = etree.SubElement(p_el, _w("r"))
    if rpr_template is not None:
        run.append(deepcopy(rpr_template))
    t_el = etree.SubElement(run, _w("t"))
    t_el.set(XML_SPACE, "preserve")
    t_el.text = text


def _append_break(p_el: etree._Element, rpr_template: etree._Element | None) -> None:
    """Append a soft line break (``<w:br/>``) carrying the same run formatting."""
    run = etree.SubElement(p_el, _w("r"))
    if rpr_template is not None:
        run.append(deepcopy(rpr_template))
    etree.SubElement(run, _w("br"))


def _append_text_with_breaks(
    p_el: etree._Element,
    text: str,
    rpr_template: etree._Element | None,
) -> None:
    """Append text where literal ``\\n`` characters become Word soft breaks."""
    if not text:
        return
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            _append_break(p_el, rpr_template)
        if line:
            _append_text_run(p_el, line, rpr_template)


def _make_body_paragraph(
    text: str,
    p_pr: etree._Element | None,
    rpr_template: etree._Element | None,
) -> etree._Element:
    """Build a ``<w:p>`` that inherits paragraph + run formatting from the anchor."""
    new_p = etree.Element(_w("p"))
    if p_pr is not None:
        new_p.insert(0, deepcopy(p_pr))
    _append_text_with_breaks(new_p, text, rpr_template)
    return new_p


def _replace_docx_internal(path: Path, internal: str, data: bytes) -> None:
    tmp_path = path.with_name(path.name + ".tmp")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            payload = data if info.filename == internal else zin.read(info.filename)
            zout.writestr(info, payload)
    tmp_path.replace(path)


def _patch_cover_letter_document_xml(document_xml: bytes, body_parts: list[str]) -> bytes:
    root = etree.fromstring(document_xml)
    body_el = root.find(_w("body"))
    if body_el is None:
        raise RuntimeError("Invalid cover letter template: missing document body.")

    tag = COVER_LETTER_BODY_TAG
    anchor: etree._Element | None = None
    for p_el in _iter_document_body_paragraphs(body_el):
        if tag in _paragraph_text_from_xml(p_el):
            anchor = p_el
            break

    if anchor is None:
        raise RuntimeError(
            f"Placeholder {tag} was not found in the document body. "
            "Add it to the main letter area of your template (not only in a text box)."
        )

    full_text = _paragraph_text_from_xml(anchor)
    before, _, after = full_text.partition(tag)

    rpr_template = _extract_anchor_rpr(anchor)
    p_pr = _clear_paragraph_content(anchor)

    prefix = before.rstrip("\n")
    if prefix.strip():
        _append_text_with_breaks(anchor, prefix, rpr_template)

    cursor = anchor
    for part in body_parts:
        new_p = _make_body_paragraph(part, p_pr, rpr_template)
        cursor.addnext(new_p)
        cursor = new_p

    suffix = after.lstrip("\n")
    if suffix.strip():
        cursor.addnext(_make_body_paragraph(suffix, p_pr, rpr_template))

    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )


def fill_cover_letter_template(
    template_path: Path,
    output_path: Path,
    cover_letter_body: str,
) -> Path:
    """Fill the cover letter template with AI-generated body text only.

    Copies the template byte-for-byte except ``word/document.xml``, so headers,
    footers, drawings, and styles are preserved. Only ``{{COVER_LETTER_BODY}}`` is
    replaced - including when Word split the placeholder across multiple runs.

    The inserted paragraphs inherit the anchor paragraph's run formatting
    (``<w:rPr>``: font family, size, color, language) so the body uses the same
    typography the user designed in their template. Single ``\\n`` characters
    inside a paragraph are emitted as ``<w:br/>`` soft line breaks (used by the
    sign-off block: ``Best regards,\\nFull Name``).
    """
    body_parts = [p.strip("\r ") for p in cover_letter_body.split("\n\n") if p.strip()]
    if not body_parts:
        raise RuntimeError("Cover letter body is empty - nothing to insert into the template.")

    shutil.copy2(template_path, output_path)
    with zipfile.ZipFile(output_path, "r") as zf:
        document_xml = zf.read("word/document.xml")

    patched_xml = _patch_cover_letter_document_xml(document_xml, body_parts)
    _replace_docx_internal(output_path, "word/document.xml", patched_xml)
    logger.info("cover_letter_docx_created", path=str(output_path))
    return output_path


def _find_libreoffice() -> str | None:
    settings = get_settings()
    if settings.libreoffice_path:
        p = Path(settings.libreoffice_path)
        if p.exists():
            return str(p)

    candidates = [shutil.which("soffice"), shutil.which("libreoffice")]
    if platform.system() == "Windows":
        candidates.extend([
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ])
    for c in candidates:
        if c and Path(c).exists():
            return str(c)
    return None


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Convert DOCX to PDF using LibreOffice headless."""
    libre = _find_libreoffice()
    if not libre:
        raise RuntimeError("LibreOffice not found. Install LibreOffice or set LIBREOFFICE_PATH.")

    if pdf_path.exists():
        pdf_path.unlink()

    outdir = pdf_path.parent
    cmd = [libre, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)]

    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=120)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LibreOffice conversion failed: {e.stderr}") from e
    except subprocess.TimeoutExpired as e:
        raise RuntimeError("LibreOffice conversion timed out after 120s") from e

    produced = outdir / f"{docx_path.stem}.pdf"
    if not produced.exists():
        raise RuntimeError("LibreOffice finished but PDF was not produced")

    if produced.resolve() != pdf_path.resolve():
        if pdf_path.exists():
            pdf_path.unlink()
        produced.replace(pdf_path)

    logger.info("pdf_created", path=str(pdf_path))
    return pdf_path


def build_output_directory(
    first_name: str,
    last_name: str,
    company_name: str,
    position_title: str,
) -> Path:
    """Create the output directory for resume files."""
    settings = get_settings()
    root = Path(settings.resume_output_root)

    person_dir = f"{first_name}_{last_name}".strip("_") or "Unknown"
    now = datetime.now()
    timestamp = now.strftime("%Y-%m-%d_%H-%M")
    company_clean = re.sub(r"[^\w\s-]", "", company_name or "Unknown").strip().replace(" ", "_")[:50]
    position_clean = re.sub(r"[^\w\s-]", "", position_title or "Unknown").strip().replace(" ", "_")[:50]
    job_dir = f"{timestamp}_{company_clean}_{position_clean}"

    full_path = root / person_dir / job_dir
    full_path.mkdir(parents=True, exist_ok=True)
    return full_path
