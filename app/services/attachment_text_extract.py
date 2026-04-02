"""
Extract plain text from supported attachment types for job-URL discovery.
"""

from __future__ import annotations

import io
import re
from html.parser import HTMLParser

from app.core.logging import get_logger

logger = get_logger(__name__)

_ALLOWED_EXT = frozenset(
    {
        ".txt",
        ".md",
        ".markdown",
        ".html",
        ".htm",
        ".docx",
        ".xlsx",
    }
)


class _HTMLToText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []

    def handle_data(self, data: str) -> None:
        self._chunks.append(data)

    def get_text(self) -> str:
        return re.sub(r"\s+", " ", "".join(self._chunks)).strip()


def _decode_utf8(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="replace")


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = filename.lower().rsplit("/", 1)[-1]
    if "." not in name:
        raise ValueError("File must have an extension")
    ext = "." + name.rsplit(".", 1)[-1]
    if ext not in _ALLOWED_EXT:
        raise ValueError(
            f"Unsupported file type {ext}. Allowed: {', '.join(sorted(_ALLOWED_EXT))}"
        )

    if ext in (".txt", ".md", ".markdown"):
        return _decode_utf8(data)

    if ext in (".html", ".htm"):
        parser = _HTMLToText()
        parser.feed(_decode_utf8(data))
        return parser.get_text() or _decode_utf8(data)

    if ext == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)

    if ext == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None and str(cell).strip():
                        lines.append(str(cell).strip())
        wb.close()
        return "\n".join(lines)

    raise ValueError(f"Unhandled extension {ext}")


def combine_file_texts(parts: list[tuple[str, str]]) -> str:
    """(filename, text) pairs separated for model context."""
    blocks: list[str] = []
    for name, text in parts:
        if not text or not text.strip():
            continue
        blocks.append(f"--- File: {name} ---\n{text.strip()}")
    return "\n\n".join(blocks)
