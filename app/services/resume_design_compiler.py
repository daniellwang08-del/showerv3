"""Compile a ResumeDesign into a styled .docx template + matching blueprint.

The generated template carries the same placeholder tags the existing fill engine
understands ({{PROFILE_SUMMARY}}, {{SKILLS_CONTENT}}, {{EXP_1..N}}), so no changes to
``resume_blueprint_renderer`` / ``resume_builder_service`` are needed - the builder
simply *produces* a template instead of requiring the user to upload one.
"""

from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from app.models.database import User
from app.models.resume_design_schemas import ResumeDesign
from app.models.resume_template_schemas import ResumeTemplateAiValidation, ResumeTemplateBlueprint
from app.services.resume_builder_service import (
    _replace_docx_internal,
    _strip_phantom_header_footer_refs,
    _w,
)
from app.services.resume_context_builder import (
    _format_period,
    _format_phone,
    _full_name,
    _profile_work_rows,
)
from app.services.resume_icons import contact_icon_png
from app.services.resume_template_service import _default_blueprint_from_tags, count_work_roles

try:  # pragma: no cover - import guard for lxml
    from lxml import etree
except Exception:  # pragma: no cover
    etree = None  # type: ignore[assignment]

import zipfile

SIDEBAR_SECTIONS = {"skills", "education", "certificates"}


def _clean_url(value: str | None) -> str:
    """Strip protocol / www / trailing slash so header links read cleanly."""
    v = (value or "").strip()
    v = re.sub(r"^https?://", "", v, flags=re.IGNORECASE)
    v = re.sub(r"^www\.", "", v, flags=re.IGNORECASE)
    return v.rstrip("/")


def _hex_to_rgb(value: str) -> RGBColor:
    v = (value or "#000000").lstrip("#")
    if len(v) == 3:
        v = "".join(ch * 2 for ch in v)
    try:
        return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))
    except Exception:
        return RGBColor(0x1F, 0x29, 0x33)


def _tint(value: str, keep: float) -> str:
    """Blend *value* toward white. ``keep`` is the fraction of the original color
    retained (0 → white, 1 → original). Used for soft header bands."""
    v = (value or "#000000").lstrip("#")
    if len(v) == 3:
        v = "".join(ch * 2 for ch in v)
    try:
        r, g, b = int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16)
    except Exception:
        return "#f1f5f9"
    r = round(255 * (1 - keep) + r * keep)
    g = round(255 * (1 - keep) + g * keep)
    b = round(255 * (1 - keep) + b * keep)
    return f"#{r:02x}{g:02x}{b:02x}"


def _set_cell_background(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tc_pr.append(shd)


def _set_cell_margins(cell, *, top: int, bottom: int, left: int, right: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _emu_to_twips(emu: int) -> int:
    return int(round(emu / 635))


def _set_row_min_height(row, height_tw: int) -> None:
    """Force a table row to be at least *height_tw* twips tall (the band)."""
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(max(0, int(height_tw))))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


# Canonical child order for <w:tblPr> per the OOXML schema (CT_TblPr).
_TBLPR_ORDER = [
    "w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual", "w:tblStyleRowBandSize",
    "w:tblStyleColBandSize", "w:tblW", "w:tblJc", "w:tblCellSpacing", "w:tblInd",
    "w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption",
    "w:tblDescription",
]


def _reorder_tblpr(tbl_pr) -> None:
    index = {qn(tag): i for i, tag in enumerate(_TBLPR_ORDER)}
    children = sorted(list(tbl_pr), key=lambda c: index.get(c.tag, 999))
    for child in children:
        tbl_pr.remove(child)
    for child in children:
        tbl_pr.append(child)


def _set_table_full_bleed(table, page_width_emu: int, left_margin_emu: int) -> None:
    """Stretch *table* across the whole page width and shift it left into the margin
    so a header band touches the page's left and right edges (no margin gap)."""
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd"):
        for el in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(el)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(_emu_to_twips(page_width_emu)))
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(-_emu_to_twips(left_margin_emu)))
    tbl_pr.append(tbl_ind)
    _reorder_tblpr(tbl_pr)


def _set_run(run, *, font: str, size_pt: float, color: RGBColor, bold: bool = False, caps: bool = False) -> None:
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    if caps:
        run.font.all_caps = True
    # Ensure the east-asian / complex-script font also maps so LibreOffice picks it up.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)


def _add_bottom_border(paragraph, color_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    borders.append(bottom)


def _add_top_border(paragraph, color_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), color_hex.lstrip("#"))
    borders.insert(0, top)


def _set_run_shading(run, fill_hex: str) -> None:
    """Fill the background behind a single run (used for badge / boxed titles)."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    rpr.append(shd)


def _set_cell_borders(cell, color_hex: str, sides: set[str], sz: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    color = color_hex.lstrip("#")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        if side in sides:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _summary_border_sides(border: str) -> tuple[set[str], int]:
    if border == "full":
        return {"top", "left", "bottom", "right"}, 8
    if border == "left":
        return {"left"}, 18
    if border == "top":
        return {"top"}, 12
    if border == "bottom":
        return {"bottom"}, 12
    if border == "x":
        return {"top", "bottom"}, 8
    return set(), 8


def _heading(container, text: str, design: ResumeDesign) -> None:
    typo = design.typography
    para = container.add_paragraph()
    para.paragraph_format.space_before = Pt(design.layout.section_gap_pt)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text.upper() if typo.uppercase_headings else text)
    _set_run(
        run,
        font=typo.font_family,
        size_pt=typo.base_font_pt * typo.heading_scale,
        color=_hex_to_rgb(design.colors.heading),
        bold=True,
        caps=typo.uppercase_headings,
    )
    if design.layout.accent_rule:
        _add_bottom_border(para, design.colors.accent)


def _body(
    container,
    text: str,
    design: ResumeDesign,
    *,
    bold: bool = False,
    color_hex: str | None = None,
    size_delta: float = 0.0,
    space_after: float = 2.0,
) -> Any:
    typo = design.typography
    para = container.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = typo.line_spacing
    run = para.add_run(text)
    _set_run(
        run,
        font=typo.font_family,
        size_pt=typo.base_font_pt + size_delta,
        color=_hex_to_rgb(color_hex or design.colors.text),
        bold=bold,
    )
    return para


def _placeholder_paragraph(container, tag: str, design: ResumeDesign) -> None:
    """A body-styled paragraph holding only a placeholder tag for the fill engine."""
    _body(container, tag, design, space_after=design.layout.section_gap_pt / 2)


def _remove_leading_empty(container) -> None:
    paras = container.paragraphs
    if paras and not paras[0].runs and not paras[0].text.strip():
        el = paras[0]._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


# ── Section renderers ──────────────────────────────────────────────────────

def _render_header(
    container,
    design: ResumeDesign,
    profile: dict[str, Any],
    *,
    on_dark: bool = False,
    trailing_space_pt: float | None = None,
) -> None:
    typo = design.typography
    align = WD_ALIGN_PARAGRAPH.CENTER if design.layout.header_align == "center" else WD_ALIGN_PARAGRAPH.LEFT

    name_color = _hex_to_rgb("#ffffff" if on_dark else design.colors.heading)
    title_color = _hex_to_rgb("#f1f5f9" if on_dark else design.colors.accent)
    contact_color = _hex_to_rgb("#dbe4f0" if on_dark else design.colors.muted)

    name_para = container.add_paragraph()
    name_para.alignment = align
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(profile.get("full_name") or "Your Name")
    _set_run(
        name_run,
        font=typo.font_family,
        size_pt=typo.base_font_pt * typo.name_scale,
        color=name_color,
        bold=True,
    )
    last_para = name_para

    if profile.get("title"):
        title_para = container.add_paragraph()
        title_para.alignment = align
        title_para.paragraph_format.space_after = Pt(3)
        title_run = title_para.add_run(profile["title"])
        _set_run(
            title_run,
            font=typo.font_family,
            size_pt=typo.base_font_pt * 1.1,
            color=title_color,
            bold=False,
        )
        last_para = title_para

    contact_items = [
        (kind, val)
        for kind, val in (
            ("email", profile.get("email")),
            ("phone", profile.get("phone")),
            ("linkedin", _clean_url(profile.get("linkedin"))),
            ("github", _clean_url(profile.get("github"))),
        )
        if val
    ]
    contact_hex = "#dbe4f0" if on_dark else design.colors.muted
    contact_size = typo.base_font_pt * 0.95
    icon_style = getattr(design.layout, "contact_icons", "brand")

    def emit_item(p, kind: str, text: str) -> None:
        png = None
        if icon_style != "none":
            try:
                png = contact_icon_png(kind, contact_hex, icon_style)
            except Exception:
                png = None
        if png:
            icon_run = p.add_run()
            try:
                icon_run.add_picture(BytesIO(png), height=Pt(contact_size))
            except Exception:
                pass
            sp = p.add_run("\u2009")  # thin space between icon and text
            _set_run(sp, font=typo.font_family, size_pt=contact_size, color=contact_color)
        r = p.add_run(text)
        _set_run(r, font=typo.font_family, size_pt=contact_size, color=contact_color)

    if contact_items:
        if design.layout.contact_layout == "stacked":
            for kind, val in contact_items:
                p = container.add_paragraph()
                p.alignment = align
                p.paragraph_format.space_after = Pt(1)
                emit_item(p, kind, val)
                last_para = p
        else:
            p = container.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(2)
            for idx, (kind, val) in enumerate(contact_items):
                if idx:
                    gap = p.add_run("    ")
                    _set_run(gap, font=typo.font_family, size_pt=contact_size, color=contact_color)
                emit_item(p, kind, val)
            last_para = p

    if trailing_space_pt is not None:
        last_para.paragraph_format.space_after = Pt(trailing_space_pt)


def _decode_data_url(data_url: str) -> bytes | None:
    """Decode a ``data:image/...;base64,...`` URL into raw bytes."""
    try:
        if "," not in data_url:
            return None
        head, b64 = data_url.split(",", 1)
        if "base64" not in head:
            return None
        return base64.b64decode(b64)
    except Exception:
        return None


def _header_band_height_pt(design: ResumeDesign, profile: dict[str, Any]) -> float:
    """Estimate the header band height (pt) from typography + padding + present
    content, so the behind-text image is sized to fill exactly that band."""
    typo = design.typography
    base = typo.base_font_pt
    h = design.layout.hp_top + design.layout.hp_bottom + base * typo.name_scale * 1.18
    if profile.get("title"):
        h += base * 1.1 * 1.4
    has_contact = any(
        profile.get(k) for k in ("email", "phone", "linkedin", "github")
    )
    if has_contact:
        if design.layout.contact_layout == "stacked":
            n = sum(1 for k in ("email", "phone", "linkedin", "github") if profile.get(k))
            h += base * 0.95 * 1.55 * max(1, n)
        else:
            h += base * 0.95 * 1.55
    return h + 4


def _add_band_background_image(paragraph, image_bytes: bytes, width_emu: int, height_emu: int) -> bool:
    """Anchor *image_bytes* as a full-bleed, behind-text picture at the page's
    top-left corner. Returns True on success."""
    try:
        run = paragraph.add_run()
        run.add_picture(BytesIO(image_bytes), width=Emu(int(width_emu)), height=Emu(int(height_emu)))
        drawing = run._r.find(qn("w:drawing"))
        if drawing is None:
            return False
        inline = drawing.find(qn("wp:inline"))
        if inline is None:
            return False
        extent = inline.find(qn("wp:extent"))
        graphic = inline.find(qn("a:graphic"))
        doc_pr = inline.find(qn("wp:docPr"))
        if extent is None or graphic is None or doc_pr is None:
            return False
        cx, cy = extent.get("cx"), extent.get("cy")

        anchor = OxmlElement("wp:anchor")
        for attr, val in (
            ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
            ("simplePos", "0"), ("relativeHeight", "0"), ("behindDoc", "1"),
            ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1"),
        ):
            anchor.set(attr, val)

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        anchor.append(simple_pos)

        for tag, rel in (("wp:positionH", "page"), ("wp:positionV", "page")):
            pos = OxmlElement(tag)
            pos.set("relativeFrom", rel)
            off = OxmlElement("wp:posOffset")
            off.text = "0"
            pos.append(off)
            anchor.append(pos)

        ext = OxmlElement("wp:extent")
        ext.set("cx", cx)
        ext.set("cy", cy)
        anchor.append(ext)

        effect = OxmlElement("wp:effectExtent")
        for a in ("l", "t", "r", "b"):
            effect.set(a, "0")
        anchor.append(effect)

        anchor.append(OxmlElement("wp:wrapNone"))
        anchor.append(doc_pr)
        anchor.append(OxmlElement("wp:cNvGraphicFramePr"))
        anchor.append(graphic)

        drawing.remove(inline)
        drawing.append(anchor)
        return True
    except Exception:
        return False


def _render_header_band(doc, design: ResumeDesign, profile: dict[str, Any], section) -> None:
    """Render the header inside a full-bleed shaded band (single-cell table) that
    touches the page's top, left, and right edges. The section top margin is set to
    zero by the caller; cell padding keeps the text aligned with the body margins.

    When ``header_background == "image"`` the band fill is replaced by a behind-text
    picture anchored to the page, sized to the estimated band height."""
    bg = design.layout.header_background
    header_image = design.layout.header_image if bg == "image" else None
    image_bytes = _decode_data_url(header_image.data_url) if header_image else None

    if header_image and image_bytes:
        light_text = header_image.text == "light"
        on_dark = light_text
        # Fallback fill (shown only if the image fails to render) matches the text mode.
        fill = "0f172a" if light_text else "e2e8f0"
    elif bg == "solid":
        fill = (design.colors.accent or "#2563eb").lstrip("#")
        on_dark = True
    else:
        fill = _tint(design.colors.accent, 0.14).lstrip("#")
        on_dark = False

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_table_full_bleed(table, section.page_width, section.left_margin)
    cell.width = section.page_width
    _set_cell_background(cell, fill)
    # Each band padding side is independently controllable; 1 pt = 20 twips.
    lay = design.layout
    _set_cell_margins(
        cell,
        top=max(0, int(round(lay.hp_top * 20))),
        bottom=max(0, int(round(lay.hp_bottom * 20))),
        left=max(0, int(round(lay.hp_left * 20))),
        right=max(0, int(round(lay.hp_right * 20))),
    )
    _render_header(cell, design, profile, on_dark=on_dark)
    _remove_leading_empty(cell)

    if header_image and image_bytes:
        band_h_pt = _header_band_height_pt(design, profile)
        # Keep the row at least the image height so text and picture line up.
        _set_row_min_height(table.rows[0], int(round(band_h_pt * 20)))
        _add_band_background_image(
            cell.paragraphs[0],
            image_bytes,
            width_emu=section.page_width,
            height_emu=int(round(band_h_pt * 12700)),
        )

    # Breathing room between the band and the first section.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(design.layout.m_top * 0.4)
    spacer.paragraph_format.space_after = Pt(2)


_SUMMARY_TITLE = "Professional Summary"
_SUMMARY_TAG = "{{PROFILE_SUMMARY}}"


def _summary_emit_title(target, design: ResumeDesign, st, on_solid: bool) -> list:
    """Emit the styled summary heading into *target*. Returns the paragraphs created
    (empty for hidden / inline titles, which fold the label into the body)."""
    typo = design.typography
    # 'side' degrades to an above-title in the .docx (clean and width-safe).
    mode = "above" if st.title == "side" else st.title
    if mode in ("hidden", "inline"):
        return []

    upper = typo.uppercase_headings
    title_color = _hex_to_rgb("#ffffff" if on_solid else design.colors.heading)
    accent_color = _hex_to_rgb("#ffffff" if on_solid else design.colors.accent)
    accent_hex = "#ffffff" if on_solid else design.colors.accent
    size = typo.base_font_pt * typo.heading_scale

    p = target.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True

    if mode == "overline":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if st.align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        _add_top_border(p, accent_hex)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(_SUMMARY_TITLE.upper())
        _set_run(run, font=typo.font_family, size_pt=typo.base_font_pt * 0.95, color=title_color, bold=True, caps=True)
        return [p]

    if mode == "badge":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if st.align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        badge_bg = "#ffffff" if on_solid else design.colors.accent
        badge_fg = design.colors.accent if on_solid else "#ffffff"
        run = p.add_run(f"  {_SUMMARY_TITLE.upper()}  ")
        _set_run(run, font=typo.font_family, size_pt=typo.base_font_pt * 0.85, color=_hex_to_rgb(badge_fg), bold=True, caps=True)
        _set_run_shading(run, badge_bg)
        return [p]

    # Standard above / centered title.
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if mode == "centered" else WD_ALIGN_PARAGRAPH.LEFT
    if st.title_accent == "bar":
        bar = p.add_run("\u258f ")
        _set_run(bar, font=typo.font_family, size_pt=size, color=accent_color, bold=True)
    elif st.title_accent == "dot":
        dot = p.add_run("\u25cf ")
        _set_run(dot, font=typo.font_family, size_pt=typo.base_font_pt, color=accent_color, bold=True)

    run = p.add_run(_SUMMARY_TITLE.upper() if upper else _SUMMARY_TITLE)
    if st.title_accent == "box":
        _set_run(run, font=typo.font_family, size_pt=size, color=_hex_to_rgb("#ffffff"), bold=True, caps=upper)
        _set_run_shading(run, design.colors.accent)
    else:
        _set_run(run, font=typo.font_family, size_pt=size, color=title_color, bold=True, caps=upper)

    # 'none' accent inherits the global accent-rule so the default style matches the
    # other section headings; 'underline' always draws it.
    if st.title_accent == "underline" or (st.title_accent == "none" and design.layout.accent_rule):
        _add_bottom_border(p, accent_hex)
    return [p]


def _summary_emit_body(target, design: ResumeDesign, st, on_solid: bool):
    typo = design.typography
    body_color = _hex_to_rgb("#ffffff" if on_solid else design.colors.text)
    p = target.add_paragraph()
    p.alignment = _ALIGN_MAP.get(st.align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.line_spacing = typo.line_spacing
    p.paragraph_format.space_after = Pt(design.layout.section_gap_pt / 2)
    if st.title == "inline":
        upper = typo.uppercase_headings
        lead = p.add_run((_SUMMARY_TITLE.upper() if upper else _SUMMARY_TITLE) + ".  ")
        _set_run(
            lead,
            font=typo.font_family,
            size_pt=typo.base_font_pt,
            color=_hex_to_rgb("#ffffff" if on_solid else design.colors.heading),
            bold=True,
            caps=upper,
        )
    run = p.add_run(_SUMMARY_TAG)
    _set_run(run, font=typo.font_family, size_pt=typo.base_font_pt, color=body_color, bold=False)
    if st.italic:
        run.font.italic = True
    return p


def _render_summary(container, design: ResumeDesign) -> None:
    st = design.sections.summary_style
    on_solid = st.surface in ("solid", "gradient")
    has_box = st.surface != "none" or st.border != "none"
    gap = design.layout.section_gap_pt

    if has_box:
        # Spacer paragraph carries the section gap (tables cannot set space-before).
        spacer = container.add_paragraph()
        spacer.paragraph_format.space_before = Pt(gap)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(2)

        table = container.add_table(rows=1, cols=1)
        table.allow_autofit = False
        cell = table.rows[0].cells[0]

        if st.surface == "tint":
            _set_cell_background(cell, _tint(design.colors.accent, 0.12))
        elif st.surface in ("solid", "gradient"):
            # No real gradient in WordprocessingML; a solid accent keeps white text legible.
            _set_cell_background(cell, design.colors.accent)
        sides, sz = _summary_border_sides(st.border)
        if sides:
            _set_cell_borders(cell, "#ffffff" if on_solid else design.colors.accent, sides, sz)
        pad_tw = max(60, int(round(st.pad_pt * 20)))
        _set_cell_margins(cell, top=pad_tw, bottom=pad_tw, left=pad_tw, right=pad_tw)

        _summary_emit_title(cell, design, st, on_solid)
        _summary_emit_body(cell, design, st, on_solid)
        _remove_leading_empty(cell)
        return

    paras = _summary_emit_title(container, design, st, on_solid)
    body = _summary_emit_body(container, design, st, on_solid)
    first = paras[0] if paras else body
    first.paragraph_format.space_before = Pt(gap)


def _render_skills(container, design: ResumeDesign) -> None:
    _heading(container, "Technical Skills", design)
    _placeholder_paragraph(container, "{{SKILLS_CONTENT}}", design)


def _experience_date_text(design: ResumeDesign, style, row: dict[str, Any]) -> str:
    """Date / location / type badge text for the experience header sub-line."""
    bits: list[str] = []
    if design.sections.show_period and row.get("period"):
        bits.append(row["period"])
    if design.sections.show_location and row.get("location"):
        bits.append(row["location"])
    if style.badge_style != "hidden":
        if style.show_employment_type and row.get("employment_type"):
            bits.append(str(row["employment_type"]))
        if style.show_arrangement and row.get("job_type"):
            bits.append(str(row["job_type"]))
    return "  |  ".join(bits)


def _render_experience(container, design: ResumeDesign, rows: list[dict[str, Any]], slot_count: int) -> None:
    _heading(container, "Work Experience", design)
    style = design.sections.experience_style
    font = design.typography.font_family
    head_pt = design.typography.base_font_pt + 0.5
    company_color = _hex_to_rgb(design.colors.accent if style.accent_target == "company" else design.colors.heading)
    role_color = _hex_to_rgb(design.colors.accent if style.accent_target == "role" else design.colors.text)
    date_color = _hex_to_rgb(design.colors.accent if style.accent_target == "date" else design.colors.muted)
    # Usable text width for a right-aligned tab stop (Letter = 12_192_000 EMU wide).
    usable_emu = max(int(12_192_000 - (design.layout.m_left + design.layout.m_right) * 12_700), 1_000_000)

    for i in range(1, slot_count + 1):
        row = rows[i - 1] if i - 1 < len(rows) else {}
        company = row.get("company_name") or "Company"
        title = row.get("job_title") or "Role"
        date_text = _experience_date_text(design, style, row)

        def add_company_run(p, with_role: bool) -> None:
            c_run = p.add_run(company)
            _set_run(c_run, font=font, size_pt=head_pt, color=company_color, bold=True)
            if with_role and title:
                t_run = p.add_run(f"  -  {title}")
                _set_run(t_run, font=font, size_pt=head_pt, color=role_color, bold=False)

        def add_role_line() -> None:
            if not title:
                return
            r_para = container.add_paragraph()
            r_para.paragraph_format.space_before = Pt(0)
            r_para.paragraph_format.space_after = Pt(0)
            r_para.paragraph_format.keep_with_next = True
            r_run = r_para.add_run(title)
            _set_run(r_run, font=font, size_pt=head_pt, color=role_color, bold=False)

        def add_date_subline() -> None:
            if date_text:
                _body(container, date_text, design, color_hex=design.colors.muted, size_delta=-0.5, space_after=1.0)

        header_para = container.add_paragraph()
        header_para.paragraph_format.space_before = Pt(4)
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.keep_with_next = True

        if style.header_layout == "stacked":
            # Company (line 1), role (line 2), date below or right.
            if style.date_position == "right" and date_text:
                header_para.paragraph_format.tab_stops.add_tab_stop(Emu(usable_emu), WD_TAB_ALIGNMENT.RIGHT)
                add_company_run(header_para, with_role=False)
                d_run = header_para.add_run(f"\t{date_text}")
                _set_run(d_run, font=font, size_pt=design.typography.base_font_pt - 0.5, color=date_color)
                add_role_line()
            else:
                add_company_run(header_para, with_role=False)
                add_role_line()
                add_date_subline()
        elif style.header_layout == "two_column" or style.date_position == "right":
            # Header left, date right via a right-aligned tab stop.
            if date_text:
                header_para.paragraph_format.tab_stops.add_tab_stop(Emu(usable_emu), WD_TAB_ALIGNMENT.RIGHT)
            add_company_run(header_para, with_role=True)
            if date_text:
                d_run = header_para.add_run(f"\t{date_text}")
                _set_run(d_run, font=font, size_pt=design.typography.base_font_pt - 0.5, color=date_color)
        elif style.date_position == "inline":
            add_company_run(header_para, with_role=True)
            if date_text:
                d_run = header_para.add_run(f"   \u00b7   {date_text}")
                _set_run(d_run, font=font, size_pt=design.typography.base_font_pt - 0.5, color=date_color)
        else:
            # inline header, date below
            add_company_run(header_para, with_role=True)
            add_date_subline()

        _placeholder_paragraph(container, "{{" + f"EXP_{i}" + "}}", design)


def _education_date_text(design: ResumeDesign, style, item: dict[str, Any]) -> str:
    """Period / grade / location text for the education header sub-line."""
    bits: list[str] = []
    if style.show_period and item.get("period"):
        bits.append(str(item["period"]))
    if style.show_mark and item.get("mark"):
        bits.append(str(item["mark"]))
    if style.show_location and item.get("location"):
        bits.append(str(item["location"]))
    return "  |  ".join(bits)


def _render_education(container, design: ResumeDesign, education: list[dict[str, Any]]) -> None:
    _heading(container, "Education", design)
    style = design.sections.education_style
    font = design.typography.font_family
    head_pt = design.typography.base_font_pt + 0.5
    sub_pt = design.typography.base_font_pt - 0.5
    muted = _hex_to_rgb(design.colors.muted)
    uni_color = _hex_to_rgb(design.colors.accent if style.accent_target == "university" else design.colors.heading)
    degree_color = _hex_to_rgb(design.colors.accent if style.accent_target == "degree" else design.colors.text)
    usable_emu = max(int(12_192_000 - (design.layout.m_left + design.layout.m_right) * 12_700), 1_000_000)

    for item in education:
        uni = item.get("university_name") or ""
        degree = item.get("degree") or ""
        if not (uni or degree):
            continue
        date_text = _education_date_text(design, style, item)

        def add_uni_run(p, with_degree: bool) -> None:
            if uni:
                u_run = p.add_run(uni)
                _set_run(u_run, font=font, size_pt=head_pt, color=uni_color, bold=True)
            if with_degree and degree:
                sep = "  -  " if uni else ""
                d_run = p.add_run(f"{sep}{degree}")
                _set_run(d_run, font=font, size_pt=head_pt, color=degree_color, bold=False)

        def add_degree_line() -> None:
            if not degree:
                return
            d_para = container.add_paragraph()
            d_para.paragraph_format.space_before = Pt(0)
            d_para.paragraph_format.space_after = Pt(0)
            d_para.paragraph_format.keep_with_next = True
            d_run = d_para.add_run(degree)
            _set_run(d_run, font=font, size_pt=head_pt, color=degree_color, bold=False)

        def add_date_subline() -> None:
            if date_text:
                _body(container, date_text, design, color_hex=design.colors.muted, size_delta=-0.5, space_after=1.0)

        header_para = container.add_paragraph()
        header_para.paragraph_format.space_before = Pt(4)
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.keep_with_next = True

        if style.header_layout == "stacked":
            if style.date_position == "right" and date_text:
                header_para.paragraph_format.tab_stops.add_tab_stop(Emu(usable_emu), WD_TAB_ALIGNMENT.RIGHT)
                add_uni_run(header_para, with_degree=False)
                t_run = header_para.add_run(f"\t{date_text}")
                _set_run(t_run, font=font, size_pt=sub_pt, color=muted)
                add_degree_line()
            else:
                add_uni_run(header_para, with_degree=False)
                add_degree_line()
                add_date_subline()
        elif style.date_position == "right":
            if date_text:
                header_para.paragraph_format.tab_stops.add_tab_stop(Emu(usable_emu), WD_TAB_ALIGNMENT.RIGHT)
            add_uni_run(header_para, with_degree=True)
            if date_text:
                t_run = header_para.add_run(f"\t{date_text}")
                _set_run(t_run, font=font, size_pt=sub_pt, color=muted)
        elif style.date_position == "inline":
            add_uni_run(header_para, with_degree=True)
            if date_text:
                t_run = header_para.add_run(f"   \u00b7   {date_text}")
                _set_run(t_run, font=font, size_pt=sub_pt, color=muted)
        else:  # date below
            add_uni_run(header_para, with_degree=True)
            add_date_subline()

        if style.show_description and (item.get("description") or "").strip():
            _body(container, str(item["description"]).strip(), design, space_after=2.0)


_CERT_GLYPH: dict[str, str] = {
    "dot": "\u2022",
    "dash": "\u2013",
    "check": "\u2713",
    "arrow": "\u2192",
    "square": "\u25AA",
    "none": "",
}


def _render_certificates(container, design: ResumeDesign, certificates: list[dict[str, Any]]) -> None:
    _heading(container, "Certifications", design)
    style = design.sections.certificates_style
    names = [str(item.get("name") or "").strip() for item in certificates]
    names = [n for n in names if n]
    if not names:
        return

    if style.layout == "inline":
        _body(container, ", ".join(names), design, space_after=2.0)
    elif style.layout == "pipe":
        _body(container, "  |  ".join(names), design, space_after=2.0)
    elif style.layout == "chips":
        # Word has no chip primitive; keep the names on one flowing line.
        _body(container, "    ".join(names), design, space_after=2.0)
    else:
        # list / grid: one entry per line with the chosen marker glyph.
        glyph = _CERT_GLYPH.get(style.marker, "\u2022")
        prefix = f"{glyph}  " if glyph else ""
        for n in names:
            _body(container, f"{prefix}{n}", design, space_after=1.0)


# ── Profile extraction ─────────────────────────────────────────────────────

def _profile_dict(user: User) -> dict[str, Any]:
    return {
        "full_name": _full_name(user),
        "title": (getattr(user, "profile_title", None) or "").strip(),
        "email": (getattr(user, "profile_email", None) or getattr(user, "email", None) or "").strip(),
        "phone": _format_phone(user),
        "linkedin": (getattr(user, "linkedin_url", None) or "").strip(),
        "github": (getattr(user, "github_url", None) or "").strip(),
    }


def _education_rows(user: User) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for item in getattr(user, "education", None) or []:
        if not isinstance(item, dict):
            continue
        uni = (item.get("university_name") or "").strip()
        degree = (item.get("degree") or "").strip()
        if uni or degree:
            rows.append({
                "university_name": uni,
                "degree": degree,
                "period": _format_period(item.get("period_start"), item.get("period_end")),
                "mark": (item.get("mark") or "").strip(),
                "location": (item.get("location") or "").strip(),
                "description": (item.get("description") or "").strip(),
            })
    return rows


def _certificate_rows(user: User) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for item in getattr(user, "certificates", None) or []:
        if isinstance(item, dict) and (item.get("name") or "").strip():
            rows.append({"name": item["name"].strip()})
    return rows


def _skill_count(user: User) -> int:
    count = 0
    for item in getattr(user, "technical_skills", None) or []:
        if isinstance(item, dict) and ((item.get("category") or "").strip() or (item.get("skills") or "").strip()):
            count += 1
    return count


# ── Public API ─────────────────────────────────────────────────────────────

def _strip_headers(path: Path) -> None:
    """Drop any header/footer references python-docx injected, so the generated
    template does not reserve phantom header/footer space in LibreOffice."""
    if etree is None:
        return
    try:
        with zipfile.ZipFile(path, "r") as zf:
            document_xml = zf.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile):
        return
    root = etree.fromstring(document_xml)
    _strip_phantom_header_footer_refs(root, set())
    final_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _replace_docx_internal(path, "word/document.xml", final_xml)


def compile_design(design: ResumeDesign, user: User, out_path: Path) -> tuple[list[str], ResumeTemplateBlueprint]:
    """Build a styled .docx for *design* at *out_path*; return (tags, blueprint)."""
    profile = _profile_dict(user)
    work_rows = _profile_work_rows(user)
    education = _education_rows(user)
    certificates = _certificate_rows(user)
    has_skills = _skill_count(user) > 0
    slot_count = max(len(work_rows), 1)

    order = [s for s in design.layout.section_order if s not in set(design.layout.hidden_sections)]
    # Skip data-driven sections the profile cannot fill.
    if not has_skills and "skills" in order:
        order.remove("skills")
    if not education and "education" in order:
        order.remove("education")
    if not certificates and "certificates" in order:
        order.remove("certificates")
    if "summary" not in order:
        order.insert(0, "summary")
    if "experience" not in order:
        order.append("experience")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(design.layout.m_top)
    section.bottom_margin = Pt(design.layout.m_bottom)
    section.left_margin = Pt(design.layout.m_left)
    section.right_margin = Pt(design.layout.m_right)

    normal = doc.styles["Normal"]
    normal.font.name = design.typography.font_family
    normal.font.size = Pt(design.typography.base_font_pt)
    normal.font.color.rgb = _hex_to_rgb(design.colors.text)

    def render_section(container, sec: str) -> None:
        if sec == "summary":
            _render_summary(container, design)
        elif sec == "skills":
            _render_skills(container, design)
        elif sec == "experience":
            _render_experience(container, design, work_rows, slot_count)
        elif sec == "education":
            _render_education(container, design, education)
        elif sec == "certificates":
            _render_certificates(container, design, certificates)

    usable = max(section.page_width - section.left_margin - section.right_margin, Pt(360))
    if design.layout.header_background != "none":
        # Full-bleed band: remove the top margin so it touches the page's top edge.
        section.top_margin = Pt(0)
        _render_header_band(doc, design, profile, section)
    else:
        _render_header(doc, design, profile, trailing_space_pt=design.layout.hp_bottom)

    if design.layout.columns == 2:
        sidebar = [s for s in order if s in SIDEBAR_SECTIONS]
        main = [s for s in order if s not in SIDEBAR_SECTIONS]
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = False
        left_cell, right_cell = table.rows[0].cells
        left_cell.width = int(usable * 0.34)
        right_cell.width = int(usable * 0.66)
        for sec in main:
            render_section(right_cell, sec)
        for sec in sidebar:
            render_section(left_cell, sec)
        _remove_leading_empty(left_cell)
        _remove_leading_empty(right_cell)
    else:
        for sec in order:
            render_section(doc, sec)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _strip_headers(out_path)

    tags: list[str] = ["{{PROFILE_SUMMARY}}"]
    if has_skills and "skills" in order:
        tags.append("{{SKILLS_CONTENT}}")
    tags.extend("{{" + f"EXP_{i}" + "}}" for i in range(1, slot_count + 1))

    profile_work_count = count_work_roles(user)
    blueprint = _default_blueprint_from_tags(tags, profile_work_count)
    blueprint.engine = "legacy_exp_n"
    blueprint.ai_validation = ResumeTemplateAiValidation(
        passed=True,
        template_type="legacy_exp_n",
        summary="Generated by the resume builder from your selected theme and layout.",
        detected_required_tags=tags,
    )
    return tags, blueprint
